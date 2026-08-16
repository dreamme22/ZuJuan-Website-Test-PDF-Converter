# -*- coding: utf-8 -*-
"""
cnvpdf.py —— PDF 一键转 TXT + DOCX（原生公式版）

原理：
  1. 用 PyMuPDF 渲染 PDF 每一页为高清 PNG；
  2. 调用 AI 视觉接口（OpenAI 兼容 chat/completions，支持图片）逐页识别，
     忠实地按题目顺序还原题干与选项，并让模型把公式用 LaTeX 包裹在 \\(...\\) 中；
  3. 生成 .txt（保留 LaTeX 可读文本）；
  4. 生成 .docx：把 LaTeX 公式经 latex2mathml + Office 的 MML2OMML.XSL
     转换为 Word 原生公式对象（OMML）内嵌到正文中。

依赖：
  pip install pymupdf python-docx latex2mathml lxml
  （需要本机装有 Microsoft Office，用于 MML2OMML.XSL）

配置（config.json，支持多 profile）：
  {
    "active": "local",
    "profiles": {
      "local":  {"base_url": "http://localhost:11434/v1", "model": "qwen3-vl:4b", "api_key": "ollama"},
      "deepseek": {"base_url": "https://api.deepseek.com/v1", "model": "deepseek-v4-pro", "api_key": "sk-xxx"}
    },
    "timeout": 600
  }

用法：
  python cnvpdf.py "试卷.pdf"
  python cnvpdf.py "试卷.pdf" --profile deepseek --out "D:\\输出"
"""

import argparse
import base64
import io
import json
import os
import re
import sys
import urllib.request
import urllib.error

import pymupdf
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

import latex2mathml.converter
from lxml import etree

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
CONFIG_PATH = os.path.join(SCRIPT_DIR, "config.json")

DEFAULT_BASE_URL = "https://api.openai.com/v1"
DEFAULT_MODEL = "gpt-4o-mini"
DEFAULT_TIMEOUT = 600

# Office 自带的 MathML -> OMML 转换表
OMML_XSL = r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL"

# 识别：公式一律用 \( ... \) 包裹的 LaTeX
VISION_PROMPT = (
    "你是试卷转写助手。请把这张试卷页面的【全部题目】内容逐题、按顺序、忠实地转写为纯文本。\n"
    "要求：\n"
    "1. 完整保留标题、学生信息栏、分节标题（如“一、单选题”“二、多选题”）。\n"
    "2. 每道题：先写题号和题干，再依次列出 A/B/C/D 选项；选项跨页或没显示完整时按实际所见转写。\n"
    "3. 所有数学表达式（公式、方程、集合、区间、函数、幂、分数、根式等）一律用 LaTeX 书写，"
    "并用括号括起来：如 \\(x^{2}\\)、\\(\\frac{x-2}{x+3} > 0\\)、\\(\\{x \\mid 0 \\leq x < 4\\}\\)、"
    "\\(A \\cap B = \\varnothing\\)、\\([0, +\\infty)\\)、\\(\\sqrt{x}\\)。\n"
    "4. 中文文字、选项前的“A．B．C．D．”等保持为普通文本，不要放进公式括号。\n"
    "5. 只输出页面上的内容本身，不要写解析、答案、注释或任何额外说明。\n"
    "6. 保持原有换行结构，每道题与每个选项各占一行。"
)

MATH_RE = re.compile(r"\\\((.*?)\\\)|\$\$(.+?)\$\$|\$(.*?)\$", re.DOTALL)


def _normalize_display(text):
    """仅将【独立成行】的 $$ ... $$ 跨行显示公式合并到一行；单行 $$...$$ 不受影响。"""
    lines = text.split("\n")
    out = []
    i = 0
    while i < len(lines):
        if lines[i].strip() == "$$":
            j = i + 1
            while j < len(lines) and lines[j].strip() != "$$":
                j += 1
            if j < len(lines):
                content = " ".join(" ".join(lines[i + 1:j]).split())
                out.append("$$" + content + "$$")
                i = j + 1
                continue
        out.append(lines[i])
        i += 1
    return "\n".join(out)


def load_config(args):
    cfg = {}
    if os.path.exists(CONFIG_PATH):
        try:
            with open(CONFIG_PATH, "r", encoding="utf-8") as f:
                cfg = json.load(f)
        except Exception as e:
            print("[警告] config.json 读取失败：", e)

    profiles = cfg.get("profiles", {})
    active = args.profile or cfg.get("active")
    acfg = profiles.get(active, {}) if active else {}

    def pick(name, env, arg, default):
        if arg:
            return arg
        if os.environ.get(env):
            return os.environ.get(env)
        if acfg.get(name):
            return acfg.get(name)
        if cfg.get(name):
            return cfg.get(name)
        return default

    num_ctx = os.environ.get("CNVPDF_NUM_CTX") or acfg.get("num_ctx") or cfg.get("num_ctx")
    return {
        "api_key": pick("api_key", "CNVPDF_API_KEY", args.api_key, ""),
        "base_url": pick("base_url", "CNVPDF_BASE_URL", args.base_url, ""),
        "model": pick("model", "CNVPDF_MODEL", args.model, ""),
        "timeout": int(os.environ.get("CNVPDF_TIMEOUT") or acfg.get("timeout") or cfg.get("timeout") or DEFAULT_TIMEOUT),
        "num_ctx": int(num_ctx) if num_ctx else None,
        "native": (acfg.get("native") or cfg.get("native")),
        "profile": active or "default",
    }


def render_pages(pdf_path):
    pages_png = []
    with pymupdf.open(pdf_path) as doc:
        n = len(doc)
        for i, page in enumerate(doc):
            pix = page.get_pixmap(dpi=200)
            buf = io.BytesIO(pix.tobytes("png"))
            pages_png.append(buf.getvalue())
            print(f"  渲染第 {i+1}/{n} 页 OK")
    return pages_png


def call_vision_api(cfg, page_png):
    b64 = base64.b64encode(page_png).decode("utf-8")
    if cfg.get("native") == "ollama":
        return _call_ollama(cfg, b64)
    return _call_openai(cfg, b64)


def _call_ollama(cfg, b64):
    """本地 Ollama 原生 /api/chat（可靠支持 options.num_ctx），非流式。"""
    url = (cfg["base_url"]).rstrip("/").replace("/v1", "") + "/api/chat"
    payload = {
        "model": cfg["model"],
        "stream": False,
        "messages": [{
            "role": "user",
            "content": VISION_PROMPT,
            "images": [b64],
        }],
    }
    if cfg.get("num_ctx"):
        payload["options"] = {"num_ctx": cfg["num_ctx"]}
    req = urllib.request.Request(
        url,
        data=json.dumps(payload).encode("utf-8"),
        headers={"Content-Type": "application/json"},
    )
    try:
        with urllib.request.urlopen(req, timeout=cfg["timeout"]) as resp:
            data = json.loads(resp.read().decode("utf-8"))
    except urllib.error.HTTPError as e:
        raise RuntimeError(f"Ollama HTTP {e.code}: {e.read().decode('utf-8','ignore')[:500]}")
    except urllib.error.URLError as e:
        raise RuntimeError(f"Ollama 连接失败: {e.reason} (请确认 Ollama 已启动)")
    if "message" in data and "content" in data["message"]:
        return data["message"]["content"]
    if "error" in data:
        raise RuntimeError(f"Ollama 错误: {data['error']}")
    raise RuntimeError(f"Ollama 返回格式异常: {json.dumps(data, ensure_ascii=False)[:500]}")


def _call_openai(cfg, b64):
    """OpenAI 兼容 chat/completions。"""
    url = (cfg["base_url"]).rstrip("/") + "/chat/completions"
    payload = {
        "model": cfg["model"],
        "messages": [{
            "role": "user",
            "content": [
                {"type": "text", "text": VISION_PROMPT},
                {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64}"}},
            ],
        }],
        "temperature": 0,
    }
    if cfg.get("num_ctx"):
        payload["options"] = {"num_ctx": cfg["num_ctx"]}
    req = urllib.request.Request(
        url,
        data=json.dumps(payload).encode("utf-8"),
        headers={"Content-Type": "application/json", "Authorization": f"Bearer {cfg['api_key']}"},
    )
    try:
        with urllib.request.urlopen(req, timeout=cfg["timeout"]) as resp:
            data = json.loads(resp.read().decode("utf-8"))
    except urllib.error.HTTPError as e:
        raise RuntimeError(f"API HTTP {e.code}: {e.read().decode('utf-8','ignore')[:500]}")
    except urllib.error.URLError as e:
        raise RuntimeError(f"API 连接失败: {e.reason} (base_url={cfg['base_url']})")
    try:
        return data["choices"][0]["message"]["content"]
    except (KeyError, IndexError):
        raise RuntimeError(f"API 返回格式异常: {json.dumps(data, ensure_ascii=False)[:500]}")


# ---------------- 原生公式(OMML)相关 ----------------
_xslt_cache = None


def _get_xslt():
    global _xslt_cache
    if _xslt_cache is None:
        if not os.path.exists(OMML_XSL):
            raise FileNotFoundError("未找到 MML2OMML.XSL，请确认本机装有 Microsoft Office：\n" + OMML_XSL)
        _xslt_cache = etree.XSLT(etree.parse(OMML_XSL))
    return _xslt_cache


def latex_to_omml(latex):
    """LaTeX -> MathML -> OMML(Word 原生公式 XML 字符串)"""
    mathml = latex2mathml.converter.convert(latex)
    mx = etree.fromstring(mathml.encode("utf-8"))
    omml = _get_xslt()(mx)
    return etree.tostring(omml, encoding="unicode")


def _add_omml(paragraph, latex):
    omml_el = etree.fromstring(latex_to_omml(latex).encode("utf-8"))
    paragraph._p.append(omml_el)


# ---------------- docx 生成 ----------------
def _set_zh(run, size=12, bold=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")


def _latex_of(m):
    return m.group(1) or m.group(2) or m.group(3)


def _add_line(doc, line, size=12, bold=False, center=False):
    """把一行文本写入段落；其中 \\(...\\) / $...$ / $$...$$ 公式转成原生 OMML。"""
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    pos = 0
    for m in MATH_RE.finditer(line):
        pre = line[pos:m.start()]
        if pre:
            r = p.add_run(pre)
            _set_zh(r, size, bold)
        latex = _latex_of(m)
        try:
            _add_omml(p, latex)
        except Exception as e:
            print("  [公式降级为文本]", latex[:40], "->", str(e)[:60])
            r = p.add_run(latex)
            _set_zh(r, size, bold)
        pos = m.end()
    tail = line[pos:]
    if tail:
        r = p.add_run(tail)
        _set_zh(r, size, bold)


def write_txt(text, txt_path):
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(text)
    print("TXT 已生成:", txt_path)


def write_docx(text, docx_path):
    doc = Document()
    text = _normalize_display(text)
    lines = [ln.rstrip() for ln in text.splitlines()]
    first = next((ln.strip() for ln in lines if ln.strip()), "")
    if first:
        _add_line(doc, first, size=16, bold=True, center=True)
        lines = [ln for ln in lines if ln.strip() != first] or []
    for ln in lines:
        s = ln.strip()
        if not s:
            continue
        if s.startswith("学校:") or "学校:___" in s:
            _add_line(doc, ln, size=11, center=True)
        elif len(s) <= 12 and s[0] in "一二三四五六七八九十" and "、" in s[:4]:
            _add_line(doc, ln, size=14, bold=True)
        else:
            _add_line(doc, ln, size=12)
    while len(doc.paragraphs) > 1 and not doc.paragraphs[-1].text.strip():
        p = doc.paragraphs[-1]._element
        p.getparent().remove(p)
    doc.save(docx_path)
    print("DOCX 已生成:", docx_path)


# ---------------- 半自动模式 ----------------
def semi_auto_fallback(pages_png, pdf_path, out_dir):
    base = os.path.splitext(os.path.basename(pdf_path))[0]
    img_dir = os.path.join(out_dir, base + "_页面图")
    os.makedirs(img_dir, exist_ok=True)
    hint_lines = []
    try:
        with pymupdf.open(pdf_path) as doc:
            for i in range(len(doc)):
                hint_lines.append(f"===== 第 {i+1} 页（文本层，可能缺少公式）=====\n{doc[i].get_text('text').strip()}")
    except Exception as e:
        print("[警告] 提取文本层失败：", e)
    for i, png in enumerate(pages_png):
        with open(os.path.join(img_dir, f"第{i+1}页.png"), "wb") as f:
            f.write(png)
    hint_path = os.path.join(out_dir, base + "_文本层提示.txt")
    with open(hint_path, "w", encoding="utf-8") as f:
        f.write("\n\n".join(hint_lines))
    print("\n[半自动] 已生成页面图:", img_dir)
    print("请让 AI 助手(如 TraeWork)读取页面图，把公式用 \\(...\\) LaTeX 包裹录入为 .txt，")
    print("再运行： python cnvpdf.py --from-txt <结果.txt>")


def from_txt(txt_path, out_dir):
    with open(txt_path, "r", encoding="utf-8") as f:
        text = f.read()
    docx_path = os.path.join(out_dir, os.path.splitext(os.path.basename(txt_path))[0] + ".docx")
    write_docx(text, docx_path)


def main():
    parser = argparse.ArgumentParser(description="PDF 一键转 TXT + DOCX（原生公式）")
    parser.add_argument("pdf", nargs="?", help="要转换的 PDF 文件路径")
    parser.add_argument("--out", help="输出目录（默认与 PDF 同目录）")
    parser.add_argument("--profile", help="使用的配置档名（config.json 里的 profiles 键）")
    parser.add_argument("--api-key", help="临时指定 API Key")
    parser.add_argument("--base-url", help="临时指定 API 地址")
    parser.add_argument("--model", help="临时指定模型")
    parser.add_argument("--from-txt", help="将已有 txt 转为 docx（跳过识别）")
    args = parser.parse_args()

    if args.from_txt:
        out = args.out or os.path.dirname(os.path.abspath(args.from_txt))
        os.makedirs(out, exist_ok=True)
        print("使用 OMML 转换表:", OMML_XSL)
        from_txt(args.from_txt, out)
        print("完成。")
        return

    if not args.pdf:
        print("用法：python cnvpdf.py <pdf路径> [--profile 配置档] [--out 输出目录]")
        sys.exit(1)

    pdf_path = os.path.abspath(args.pdf)
    if not os.path.exists(pdf_path):
        print("找不到文件：", pdf_path)
        sys.exit(1)

    out_dir = os.path.abspath(args.out) if args.out else os.path.dirname(pdf_path)
    os.makedirs(out_dir, exist_ok=True)
    base = os.path.splitext(os.path.basename(pdf_path))[0]

    cfg = load_config(args)
    print(f"使用配置档: {cfg['profile']}  |  模型: {cfg['model']}  |  接口: {cfg['base_url']}")

    # 校验 OMML 转换表
    try:
        _get_xslt()
    except FileNotFoundError as e:
        print(e)
        sys.exit(1)

    print("开始读取 PDF：", pdf_path)
    pages_png = render_pages(pdf_path)

    # 全自动识别
    parts = []
    for i, png in enumerate(pages_png):
        print(f"  AI 识别第 {i+1}/{len(pages_png)} 页…")
        parts.append(call_vision_api(cfg, png))
    full = "\n\n".join(parts).strip() + "\n"

    txt_path = os.path.join(out_dir, base + ".txt")
    docx_path = os.path.join(out_dir, base + ".docx")
    write_txt(full, txt_path)
    write_docx(full, docx_path)
    print("\n全部完成！")
    print("  TXT :", txt_path)
    print("  DOCX:", docx_path)


if __name__ == "__main__":
    main()